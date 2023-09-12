import numpy as np

from sklearn.model_selection import train_test_split
from sklearn.metrics import roc_curve, auc, classification_report
from sklearn.model_selection import StratifiedKFold


import matplotlib.pyplot as plt

import pandas as pd
from remove_intersection import remove_intersections
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score
from sklearn.discriminant_analysis import LinearDiscriminantAnalysis

from docx import Document
from docx.shared import Inches


def export_summary(result):
    model_summary_table = result.summary().tables[1].as_html()
    df_model_summary = pd.read_html(
        model_summary_table, header=0, index_col=0)[0]

    # Save the model summary table to a Word document
    document = Document()
    document.add_heading('Logistic Regression Model Summary', level=1)

    # Create a table in the Word document
    table = document.add_table(
        rows=df_model_summary.shape[0]+1, cols=df_model_summary.shape[1])

    # Add the header row to the table
    for j in range(df_model_summary.shape[-1]):
        table.cell(0, j).text = df_model_summary.columns[j]

    # Add the data rows to the table
    for i in range(df_model_summary.shape[0]):
        for j in range(df_model_summary.shape[-1]):
            table.cell(i+1, j).text = str(df_model_summary.values[i, j])

    # Save the Word document
    document.save('model_summary.docx')


def train_lda_model(data):
    bins = [0, 50, 250, 1000, float("inf")]
    frequency_range = pd.cut(data['tweet_count'], bins=bins)
    data.insert(len(data.columns) - 2, "frequency_range", frequency_range)

    X = data.iloc[:, :-3]
    y = data.iloc[:, -1]

    # Split data into training and test sets
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.3, random_state=42, stratify=data["frequency_range"])

    # Create and train the LDA model
    lda_model = LinearDiscriminantAnalysis()
    lda_model.fit(X_train, y_train)

    # Predict probabilities for test data
    y_pred = lda_model.predict_proba(X_test)[:, 1]

    # Calculate the ROC curve
    fpr, tpr, thresholds = roc_curve(y_test, y_pred)
    roc_auc = auc(fpr, tpr)

    # Plot the ROC curve
    plt.figure()
    plt.plot(fpr, tpr, color='blue', lw=2,
             label=f'ROC curve (AUC = {roc_auc:.2f})')
    plt.plot([0, 1], [0, 1], color='gray',
             linestyle='--', lw=2)  # Random classifier
    plt.xlabel('False Positive Rate (1 - Specificity)')
    plt.ylabel('True Positive Rate (Sensitivity)')
    plt.title('Receiver Operating Characteristic (ROC) Curve')
    plt.legend(loc='lower right')
    plt.grid(True)
    plt.show()

    # Convert probabilities to class predictions
    y_pred_class = (y_pred > 0.5).astype(int)

    # Calculate and print accuracy and classification report
    accuracy = (y_pred_class == y_test).mean()
    report = classification_report(y_test, y_pred_class)
    print(report)

# Call the function with your data
# train_lda_model(your_data)


def main(path_A, path_B, folder_path):
    df_A = pd.read_csv(folder_path.format(path_A),
                       encoding="utf-8", on_bad_lines='skip')
    df_B = pd.read_csv(folder_path.format(path_B),
                       encoding="utf-8", on_bad_lines='skip')

    processed_df = remove_intersections(df_A, df_B)
    df_A = processed_df[0]
    df_B = processed_df[1]

    df_A["class"] = 0
    df_B["class"] = 1

    df = pd.concat([df_A, df_B], ignore_index=True)
    df = df.drop("user_id", axis=1)
    df = df.iloc[:, np.r_[0:10, 20, 21, 22]]

    train_lda_model(df)


if __name__ == "__main__":
    folder_path = "data\\scores\\frame-axis\\aggregated\\{}"

    account_A = "external-regular#climatechange-aggregated.csv"
    account_B = "internal-regular#climatechange-aggregated.csv"

    main(account_A, account_B, folder_path)
