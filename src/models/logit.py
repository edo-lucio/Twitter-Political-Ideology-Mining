import numpy as np

from sklearn.model_selection import train_test_split
from sklearn.metrics import roc_curve, auc, classification_report, confusion_matrix
from sklearn.model_selection import StratifiedKFold


import matplotlib.pyplot as plt

import pandas as pd
from remove_intersection import remove_intersections
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score
from sklearn.preprocessing import StandardScaler
from mlxtend.feature_selection import SequentialFeatureSelector
import statsmodels.api as sm

from docx import Document
from docx.shared import Inches


def export_summary(result):
    model_summary_table = result.summary().tables[1].as_html()
    df_model_summary = pd.read_html(
        model_summary_table, header=0, index_col=0)[0]

    # Save the model summary table to a Word document
    document = Document()
    document.add_heading('Logistic Regression Model Summary', level=1)

    # Create a table in the Word document
    table = document.add_table(
        rows=df_model_summary.shape[0]+1, cols=df_model_summary.shape[1])

    # Add the header row to the table
    for j in range(df_model_summary.shape[-1]):
        table.cell(0, j).text = df_model_summary.columns[j]

    # Add the data rows to the table
    for i in range(df_model_summary.shape[0]):
        for j in range(df_model_summary.shape[-1]):
            table.cell(i+1, j).text = str(df_model_summary.values[i, j])

    # Save the Word document
    document.save('model_summary_internal_external_hashtag.docx')


def train_regression_model(data):
    # quantiles = pd.qcut(
    #     data["tweet_count"], 10, duplicates='drop', labels=[0, 1, 3, 4, 5, 6])

    # bins = [0, 50, 100, 250, 800]
    bins = [0, 50, 250, 1000, float("inf")]

    frequency_range = pd.cut(data['tweet_count'], bins=bins)

    data.insert(len(data.columns) - 2, "frequency_range", frequency_range)

    # Create design matrix X and response variable y
    X = data.iloc[:, :-3]
    y = data.iloc[:, -1]

    print(X.head())
    # add intercept term to feature matrix
    X = sm.add_constant(X)

    # split data into training and test sets
    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.3, random_state=42, stratify=data["frequency_range"])

    X_train.info()
    y_train.info()

    print(X_train)

    # create logistic regression model
    logit_model = sm.Logit(y_train, X_train.astype(float))

    # fit the model to the training data
    result = logit_model.fit()

    export_summary(result)

    # print the summary of the model
    print(result.summary())

    # evaluate the model on the test data
    y_pred = result.predict(X_test)

    # Calculate the ROC curve
    fpr, tpr, thresholds = roc_curve(y_test, y_pred)
    roc_auc = auc(fpr, tpr)

    # Plot the ROC curve
    plt.figure()
    plt.plot(fpr, tpr, color='blue', lw=2,
             label=f'ROC curve (AUC = {roc_auc:.2f})')
    plt.plot([0, 1], [0, 1], color='gray',
             linestyle='--', lw=2)  # Random classifier
    plt.xlabel('False Positive Rate (1 - Specificity)')
    plt.ylabel('True Positive Rate (Sensitivity)')
    plt.title('Receiver Operating Characteristic (ROC) Curve')
    plt.legend(loc='lower right')
    plt.grid(True)
    plt.show()

    # Convert probabilities to class predictions
    y_pred_class = (y_pred > 0.5).astype(int)

    # Calculate and print accuracy and classification report
    accuracy = (y_pred_class == y_test).mean()
    report = classification_report(y_test, y_pred_class)
    confusion_mat = confusion_matrix(y_test, y_pred_class)

    print(confusion_mat)
    print(report)


def kfold(data, n_splits=5):
    bins = [0, 50, 250, 1000, float("inf")]
    frequency_range = pd.cut(data['tweet_count'], bins=bins)
    data.insert(len(data.columns) - 2, "frequency_range", frequency_range)

    # Create design matrix X and response variable y
    X = data.iloc[:, :-3]
    y = data.iloc[:, -1]

    # add intercept term to feature matrix
    X = sm.add_constant(X)

    # Initialize k-fold cross-validation
    skf = StratifiedKFold(n_splits=n_splits, shuffle=True, random_state=42)

    # Lists to store evaluation metrics
    auc_scores = []
    f1_scores = []
    precision_scores = []

    # Perform k-fold cross-validation
    for train_index, test_index in skf.split(X, y):
        X_train, X_test = X.iloc[train_index], X.iloc[test_index]
        y_train, y_test = y.iloc[train_index], y.iloc[test_index]

        # create logistic regression model
        logit_model = sm.Logit(y_train, X_train.astype(float))

        # fit the model to the training data
        result = logit_model.fit()

        # evaluate the model on the test data
        y_pred = result.predict(X_test)

        # Convert probabilities to class predictions
        y_pred_class = (y_pred > 0.5).astype(int)

        # Calculate AUC, F1-score, and Precision
        fpr, tpr, _ = roc_curve(y_test, y_pred)
        roc_auc = auc(fpr, tpr)
        f1 = (2 * precision_score(y_test, y_pred_class) * np.mean(tpr)) / \
            (precision_score(y_test, y_pred_class) + np.mean(tpr))
        precision = precision_score(y_test, y_pred_class)

        # Append scores to lists
        auc_scores.append(roc_auc)
        f1_scores.append(f1)
        precision_scores.append(precision)

    # Create a line plot for each metric
    plt.figure(figsize=(10, 6))
    plt.plot(range(1, n_splits + 1), auc_scores, marker='o', label='AUC')
    plt.plot(range(1, n_splits + 1), f1_scores, marker='o', label='F1-score')
    plt.plot(range(1, n_splits + 1), precision_scores,
             marker='o', label='Precision')
    plt.xlabel('Fold')
    plt.ylabel('Score')
    plt.title('Evaluation Metrics Across Folds')
    plt.xticks(range(1, n_splits + 1))
    plt.legend()
    plt.grid(True)
    plt.show()


def main(path_A, path_B, folder_path):
    df_A = pd.read_csv(folder_path.format(path_A),
                       encoding="utf-8", on_bad_lines='skip')
    df_B = pd.read_csv(folder_path.format(path_B),
                       encoding="utf-8", on_bad_lines='skip')

    processed_df = remove_intersections(df_A, df_B)
    df_A = processed_df[0]
    df_B = processed_df[1]

    df_A["class"] = 0
    df_B["class"] = 1

    print(len(df_A))
    print(len(df_B))

    df = pd.concat([df_A, df_B], ignore_index=True)
    df = df.drop("user_id", axis=1)
    df = df.iloc[:, np.r_[0:10, 20, 21, 22]]

    train_regression_model(df)


if __name__ == "__main__":
    folder_path = "data\\scores\\frame-axis\\aggregated\\{}"

    account_A = "external-regular#climatechange-aggregated.csv"
    account_B = "internal-regular#climatechange-aggregated.csv"

    main(account_A, account_B, folder_path)
